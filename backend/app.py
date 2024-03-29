import requests
from werkzeug.utils import secure_filename
import os
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from textblob import TextBlob
import spacy
from docx import Document as DocxDocument
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_cors import CORS
from config import Config
from models import db, Document, Paragraph, Sentence, Keyword, User
from openai import OpenAI
from flask import request

import threading
from queue import Queue
from flask import Flask, redirect, url_for, flash, jsonify, session

from flask_login import LoginManager, login_user, logout_user, login_required, UserMixin, current_user
from flask_dance.contrib.google import make_google_blueprint, google
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# ai
client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),
)

GOOGLE_SEARCH_API_URL = "https://www.googleapis.com/customsearch/v1"
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_CX = os.getenv("GOOGLE_CX")

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
app = Flask(__name__)
app.config.from_object('config.Config')
app.config['SESSION_COOKIE_HTTPONLY'] = False
CORS(app, supports_credentials=True)
nlp = spacy.load("en_core_web_sm")
db.init_app(app)
migrate = Migrate(app, db)


file_processing_queue = Queue()

# Number of worker threads
num_worker_threads = 3  # Example: 5 threads


def process_file_from_queue():
    while True:
        task = file_processing_queue.get()
        filepath, filename, user_id = task  # Unpack userId
        with app.app_context():
            try:
                process_file(filepath, filename, user_id)  # Pass userId
            finally:
                file_processing_queue.task_done()


# Start multiple worker threads
for i in range(num_worker_threads):
    t = threading.Thread(target=process_file_from_queue, daemon=True)
    t.start()


def process_file(filepath, filename, user_id):
    text = ""
    if filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(filepath)
    elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
        text = extract_text_from_image(filepath)
    elif filename.lower().endswith('.txt'):
        text = extract_text_from_txt(filepath)
    elif filename.lower().endswith('.docx'):
        text = extract_text_from_docx(filepath)
    full_text = text

    document_sentiment = analyze_sentiment(full_text)
    document = Document(content=full_text, sentiment=document_sentiment, filename=filename, user_id=user_id)  # Use user_id
    db.session.add(document)

    paragraphs_text = split_into_paragraphs(full_text)
    for paragraph_text in paragraphs_text:
        paragraph_sentiment = analyze_sentiment(paragraph_text)
        paragraph = Paragraph(document=document, content=paragraph_text, sentiment=paragraph_sentiment)
        db.session.add(paragraph)

        sentences_text = split_into_sentences(paragraph_text)
        for sentence_text in sentences_text:
            sentence_sentiment = analyze_sentiment(sentence_text)
            sentence = Sentence(paragraph=paragraph, content=sentence_text, sentiment=sentence_sentiment)
            db.session.add(sentence)

            keywords = extract_keywords(sentence_text)
            for word in keywords:
                keyword = Keyword(sentence=sentence, word=word)
                db.session.add(keyword)

    db.session.commit()

    return jsonify({'message': f'{filename} uploaded successfully', 'sentiment': document_sentiment}), 200


threading.Thread(target=process_file_from_queue, daemon=True).start()


@app.route('/upload', methods=['POST'])
def upload_file():
    files = request.files.getlist('file')
    userId = request.form.get('userId')
    if not userId:
        response = jsonify({'error': 'UserId is required'}), 400
        return response
    if not files or all(file.filename == '' for file in files):
        response = jsonify({'error': 'No files selected'}), 400
        return response

    processed_files = []

    for file in files:
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            file_processing_queue.put((filepath, filename, userId))  # Include userId in the queue
            processed_files.append(filename)

    if not processed_files:
        response = jsonify({'error': 'No valid files were processed'}), 400
        return response

    response = jsonify({'message': f'Files queued for processing: {", ".join(processed_files)}'}), 202
    return response



@app.route('/api/documents/user/<user_id>', methods=['GET'])
def get_user_documents(user_id):
    try:
        documents = Document.query.filter_by(user_id=user_id).all()
        documents_data = [
            {'filename': doc.filename, 'documentId': doc.id, 'sentiment': doc.sentiment}
            for doc in documents
        ]
        return jsonify(documents_data), 200
    except Exception as e:
        return jsonify({'error': 'Internal Server Error', 'message': str(e)}), 500




@app.route('/document/summary', methods=['POST'])
def document_summary():
    data = request.json
    filename = data.get('filename')

    # Fetch document text from the database
    document = Document.query.filter_by(filename=filename).first()
    if not document:
        return jsonify({'error': 'Document not found'}), 404

    # Generate summary
    print(f"An error occurred")
    document_summary = get_document_summary(document.content)

    return jsonify({'filename': filename, 'summary': document_summary})


@app.route('/document/keywords', methods=['POST'])
def document_keywords():
    data = request.json
    filename = data.get('filename')  # Now using filename to identify the document

    if not filename:
        return jsonify({'error': 'Filename is required'}), 400

    # Query keywords through the relationship chain: Document -> Paragraph -> Sentence -> Keyword
    # Filter by Document.filename instead of Document.id
    keywords_query = db.session.query(Keyword.word).join(Sentence).join(Paragraph).join(Document).filter(
        Document.filename == filename).all()
    keywords = [keyword[0] for keyword in keywords_query]  # Extracting the keyword strings from the query result

    return jsonify({'keywords': keywords})


@app.route('/keyword/definition', methods=['POST'])
def keyword_definition():
    data = request.json
    keyword = data.get('keyword')

    if not keyword:
        return jsonify({'error': 'Keyword is required'}), 400

    # Now, definition should be a plain string, which is JSON serializable
    definition = get_keyword_definition(keyword)

    return jsonify({'keyword': keyword, 'definition': definition})


@app.route('/search', methods=['POST'])
def search_articles():
    keyword = request.json.get('keyword')
    if not keyword:
        return jsonify({'error': 'Keyword is required'}), 400

    params = {
        'key': GOOGLE_API_KEY,
        'cx': GOOGLE_CX,
        'q': keyword,
    }

    response = requests.get(GOOGLE_SEARCH_API_URL, params=params)
    if response.status_code == 200:
        search_results = response.json()
        links = [item['link'] for item in search_results.get('items', [])]
        return jsonify({'keyword': keyword, 'links': links})
    else:
        return jsonify({'error': 'Failed to fetch search results'}), response.status_code

@app.route('/api/filter/sentiment/<sentiment>', methods=['GET'])
def filter_by_sentiment(sentiment):
    try:
        paragraphs = Paragraph.query.filter_by(sentiment=sentiment).all()
        sentences = Sentence.query.filter_by(sentiment=sentiment).all()

        paragraphs_data = [{'id': p.id, 'content': p.content, 'sentiment': p.sentiment} for p in paragraphs]
        sentences_data = [{'id': s.id, 'content': s.content, 'sentiment': s.sentiment} for s in sentences]

        return jsonify({'paragraphs': paragraphs_data, 'sentences': sentences_data}), 200
    except Exception as e:
        return jsonify({'error': 'Internal Server Error', 'message': str(e)}), 500

@app.route('/api/search/keyword', methods=['POST'])
def search_by_keyword():
    keyword = request.json.get('keyword')
    if not keyword:
        return jsonify({'error': 'Keyword is required'}), 400

    sentences = Sentence.query.join(Sentence.keywords).filter(Keyword.word == keyword).all()
    sentence_ids = [sentence.id for sentence in sentences]

    paragraphs = Paragraph.query.join(Paragraph.sentences).filter(Sentence.id.in_(sentence_ids)).all()

    sentences_data = [{'id': sentence.id, 'content': sentence.content} for sentence in sentences]
    paragraphs_data = [{'id': paragraph.id, 'content': paragraph.content} for paragraph in paragraphs]

    return jsonify({'sentences': sentences_data, 'paragraphs': paragraphs_data}), 200



def get_document_summary(document_text):
    try:
        completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are an intelligent assistant."},
                {"role": "user", "content": f"Summarize the following document:\n\n{document_text}"}
            ],
            model="gpt-3.5-turbo"
        )
        return completion.choices[0].message.content
    except Exception as e:
        # Handle exceptions or log error and return an informative message or raise the error
        print(f"An error occurred: {str(e)}")
        raise

def get_keyword_definition(keyword):
    try:
        completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are an intelligent assistant."},
                {"role": "user", "content": f"Define the word: {keyword}."}
            ],
            model="gpt-3.5-turbo"
        )
        return completion.choices[0].message.content
    except Exception as e:
        # Handle exceptions or log error and return an informative message or raise the error
        print(f"An error occurred: {str(e)}")
        raise


def analyze_sentiment(text):
    """Simple sentiment analysis using TextBlob."""
    blob = TextBlob(text)
    return 'positive' if blob.sentiment.polarity > 0 else 'negative' if blob.sentiment.polarity < 0 else 'neutral'


def split_into_paragraphs(text):
    """Split text into paragraphs based on newline characters."""
    return text.split('\n\n')


def split_into_sentences(paragraph):
    """Split paragraph into sentences using spaCy."""
    doc = nlp(paragraph)
    return [sent.text for sent in doc.sents]


def extract_keywords(sentence):
    """Extract keywords from a sentence using spaCy."""
    doc = nlp(sentence)
    return [token.lemma_ for token in doc if token.pos_ in ['NOUN', 'PROPN']]


def extract_text_from_pdf(filepath):
    text = ""
    with fitz.open(filepath) as doc:
        for page in doc:
            text += page.get_text()
    return text


def extract_text_from_image(filepath):
    return pytesseract.image_to_string(Image.open(filepath))


def extract_text_from_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as file:
        text = file.read()
    return text


def extract_text_from_docx(filepath):
    doc = DocxDocument(filepath)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


if __name__ == '__main__':
    with app.app_context():
        db.create_all()  # Ensure database tables are created
    app.run(debug=True)
