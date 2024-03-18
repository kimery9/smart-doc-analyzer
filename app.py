import requests
from flask import Flask, request, jsonify, session
from werkzeug.utils import secure_filename
import os
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from textblob import TextBlob
import spacy
from docx import Document as DocxDocument
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_cors import CORS
from config import Config
from models import db, Document, Paragraph, Sentence, Keyword, User
from openai import OpenAI
from flask import render_template, redirect, url_for, request, flash
from flask_login import login_user, logout_user, current_user, login_required, LoginManager, UserMixin

#ai
client = OpenAI()

GOOGLE_SEARCH_API_URL = "https://www.googleapis.com/customsearch/v1"
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_CX = os.getenv("GOOGLE_CX")

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
app = Flask(__name__)
app.config.from_object('config.Config')
app.config['SESSION_COOKIE_HTTPONLY'] = False
CORS(app, supports_credentials=True)
db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = Config.LOGIN_VIEW

# Load spaCy model
nlp = spacy.load("en_core_web_sm")


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


@app.route('/login', methods=['GET', 'POST'])
def login():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    user = User.query.filter_by(username=username).first()
    if user and user.check_password(password):
        login_user(user, remember=True)
        session['user_id'] = user.id
        # Query for the user's documents
        documents = Document.query.filter_by(user_id=user.id).all()
        # Format documents for JSON response
        documents_data = [{'filename': doc.filename, 'id': doc.id} for doc in documents]
        return jsonify({'message': 'Logged in successfully', 'documents': documents_data}), 200
    return jsonify({'error': 'Invalid username or password'}), 401


@app.route('/logout', methods=['POST'])
@login_required
def logout():
    logout_user()
    return jsonify({'message': 'Logged out successfully'}), 200


@app.route('/register', methods=['POST'])
def register():
    data = request.json
    username = data.get('username')
    password = data.get('password')

    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username already exists'}), 400

    new_user = User(username=username)
    new_user.set_password(password)  # Assuming set_password method exists

    db.session.add(new_user)
    db.session.commit()

    return jsonify({'message': 'User registered successfully'}), 201



@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    text = ""
    if filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(filepath)
    elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
        text = extract_text_from_image(filepath)
    elif filename.lower().endswith('.txt'):
        text = extract_text_from_txt(filepath)
    elif filename.lower().endswith('.docx'):
        text = extract_text_from_docx(filepath)
    else:
        return jsonify({'error': 'Unsupported file type'}), 400

    full_text = text

    document_sentiment = analyze_sentiment(full_text)
    document = Document(content=full_text, sentiment=document_sentiment, filename=filename, user_id='1')
    db.session.add(document)

    paragraphs_text = split_into_paragraphs(full_text)
    for paragraph_text in paragraphs_text:
        paragraph_sentiment = analyze_sentiment(paragraph_text)
        paragraph = Paragraph(document=document, content=paragraph_text, sentiment=paragraph_sentiment)
        db.session.add(paragraph)

        sentences_text = split_into_sentences(paragraph_text)
        for sentence_text in sentences_text:
            sentence_sentiment = analyze_sentiment(sentence_text)
            sentence = Sentence(paragraph=paragraph, content=sentence_text, sentiment=sentence_sentiment)
            db.session.add(sentence)

            keywords = extract_keywords(sentence_text)
            for word in keywords:
                keyword = Keyword(sentence=sentence, word=word)
                db.session.add(keyword)

    db.session.commit()

    return jsonify({'message': f'{filename} uploaded successfully', 'document_sentiment': document_sentiment}), 200


@app.route('/document/summary', methods=['POST'])
def document_summary():
    data = request.json
    filename = data.get('filename')

    # Fetch document text from the database
    document = Document.query.filter_by(filename=filename).first()
    if not document:
        return jsonify({'error': 'Document not found'}), 404

    # Generate summary
    document_summary = get_document_summary(document.content)

    return jsonify({'filename': filename, 'summary': document_summary})

@app.route('/document/keywords', methods=['POST'])
def document_keywords():
    data = request.json
    filename = data.get('filename')  # Now using filename to identify the document

    if not filename:
        return jsonify({'error': 'Filename is required'}), 400

    # Query keywords through the relationship chain: Document -> Paragraph -> Sentence -> Keyword
    # Filter by Document.filename instead of Document.id
    keywords_query = db.session.query(Keyword.word).join(Sentence).join(Paragraph).join(Document).filter(Document.filename == filename).all()
    keywords = [keyword[0] for keyword in keywords_query]  # Extracting the keyword strings from the query result

    return jsonify({'keywords': keywords})



@app.route('/keyword/definition', methods=['POST'])
def keyword_definition():
    data = request.json
    keyword = data.get('keyword')

    if not keyword:
        return jsonify({'error': 'Keyword is required'}), 400

    # Now, definition should be a plain string, which is JSON serializable
    definition = get_keyword_definition(keyword)

    return jsonify({'keyword': keyword, 'definition': definition})


@app.route('/search', methods=['POST'])
def search_articles():
    keyword = request.json.get('keyword')
    if not keyword:
        return jsonify({'error': 'Keyword is required'}), 400

    params = {
        'key': GOOGLE_API_KEY,
        'cx': GOOGLE_CX,
        'q': keyword,
    }

    response = requests.get(GOOGLE_SEARCH_API_URL, params=params)
    if response.status_code == 200:
        search_results = response.json()
        links = [item['link'] for item in search_results.get('items', [])]
        return jsonify({'keyword': keyword, 'links': links})
    else:
        return jsonify({'error': 'Failed to fetch search results'}), response.status_code


def get_keyword_definition(keyword):
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an intelligent assistant."},
            {"role": "user", "content": f"Define the word: {keyword}."}
        ]
    )

    return completion.choices[0].message.content


def get_document_summary(document_text):
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an intelligent assistant."},
            {"role": "user", "content": f"Summarize the following document:\n\n{document_text}"}
        ]
    )

    return completion.choices[0].message.content


def analyze_sentiment(text):
    """Simple sentiment analysis using TextBlob."""
    blob = TextBlob(text)
    return 'positive' if blob.sentiment.polarity > 0 else 'negative' if blob.sentiment.polarity < 0 else 'neutral'


def split_into_paragraphs(text):
    """Split text into paragraphs based on newline characters."""
    return text.split('\n\n')


def split_into_sentences(paragraph):
    """Split paragraph into sentences using spaCy."""
    doc = nlp(paragraph)
    return [sent.text for sent in doc.sents]


def extract_keywords(sentence):
    """Extract keywords from a sentence using spaCy."""
    doc = nlp(sentence)
    return [token.lemma_ for token in doc if token.pos_ in ['NOUN', 'PROPN']]


def extract_text_from_pdf(filepath):
    text = ""
    with fitz.open(filepath) as doc:
        for page in doc:
            text += page.get_text()
    return text


def extract_text_from_image(filepath):
    return pytesseract.image_to_string(Image.open(filepath))


def extract_text_from_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as file:
        text = file.read()
    return text


def extract_text_from_docx(filepath):
    doc = DocxDocument(filepath)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


with app.app_context():
    db.create_all()  # Ensure database tables are created
app.run(debug=True)
